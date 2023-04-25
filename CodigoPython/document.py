from docx import Document
from PySimpleGUI import PySimpleGUI as sg
import translate
from translate import Translator
import datetime
import pandas as pd
import win32com.client as win32

traducao = Translator(from_lang="english", to_lang="pt-br")
datatempo = datetime.date.today()

#Função
def preencherDados():
    if value['Entrega'] == True:
        documento = Document('.\Termos\Termo de Entrega.docx')
    if value['Emprestimo'] == True:
        documento = Document('.\Termos\Termo de Emprestimo.docx')
    if value['Devolucao'] == True:
        documento = Document('.\Termos\Termo de Devolucao.docx')

    referencias = {
        'Chamado': value['chamado'],
        'Pessoa': value['nome'],
        'DD': datatempo.strftime("%d"),
        'MM': traducao.translate(datatempo.strftime("%B ")),
        'YY': datatempo.strftime("%Y "),
    }
    for paragrafo in documento.paragraphs:
        for codigo in referencias:
            if 'Chamado' in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(codigo,referencias[codigo])
            if 'Pessoa' in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(codigo,referencias[codigo])
            if 'DD' in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(codigo,referencias[codigo])
            if 'MM' in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(codigo,referencias[codigo])
            if 'YY' in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(codigo,referencias[codigo])
            
        documento.tables
        documento.tables[0].cell(1,0).text = 'Nome: ',value['nome']
        documento.tables[0].cell(1,1).text = 'Área / Núcleo: ', value['area'], ' - 305'
        documento.tables[1].cell(2,1).text = value['patrimonio']
        documento.tables[1].cell(2,2).text = value['descricao']
        documento.tables[2].cell(1,0).text = value['observacoes']

    if value['Notebook'] == True:
        nomeDocumentos = 'TE - ' + str(value['nome']) + ' - Notebook.docx'
        if value['Devolucao'] == True:
            nomeDocumentos = nomeDocumentos.replace('TE', 'TD')
        if value['Emprestimo'] == True:
            nomeDocumentos = nomeDocumentos.replace('TE', 'TEM')
        documento.save(nomeDocumentos)
    if value['Desktop'] == True:
        nomeDocumentos = 'TE - ' + str(value['nome']) + ' - Desktop.docx'
        if value['Devolucao'] == True:
            nomeDocumentos = nomeDocumentos.replace('TE', 'TD')
        if value['Emprestimo'] == True:
            nomeDocumentos = nomeDocumentos.replace('TE', 'TEM')
        documento.save(nomeDocumentos)
    if value['Monitor'] == True:
        nomeDocumentos = 'TE - ' + str(value['nome']) + ' - Monitor.docx'
        if value['Devolucao'] == True:
            nomeDocumentos = nomeDocumentos.replace('TE', 'TD')
        if value['Emprestimo'] == True:
            nomeDocumentos = nomeDocumentos.replace('TE', 'TEM')
        documento.save(nomeDocumentos)

#Layout
sg.theme('LightGrey2')
layout = [
    [
        sg.Text('Tipo de termo'),
        sg.Radio('Termo de Entrega', 'group 1', default=True, key='Entrega'), 
        sg.Radio('Termo de Emprestimo', 'group 1', default=False, key='Emprestimo'), 
        sg.Radio('Termo de Devolução', 'group 1', default=False, key='Devolucao')
    ],
    [
        sg.Text('Tipo de ativo'),
        sg.Radio('Notebook', 'group 2', default=True, key='Notebook'), 
        sg.Radio('Desktop', 'group 2', default=False, key='Desktop'), 
        sg.Radio('Monitor', 'group 2', default=False, key='Monitor')
    ],
    [sg.Text('Nome da Pessoa'), sg.Input(key='nome')],
    [sg.Text('Nome da Área'), sg.Input(key='area')],
    [sg.Text('Número do Patrimonio'), sg.Input(key='patrimonio')],
    [sg.Text('Descrição'), sg.Input(key='descricao')],
    [sg.Text('Observações'), sg.Input(key='observacoes')],
    [sg.Text('Número do chamado'), sg.Input(key='chamado')],
    [sg.Button('Gerar')]
]

#Window
window = sg.Window('Termos', layout, finalize=True)
window.bind("<Return>", "_Enter")
while True:
    events, value = window.read()
    if events == sg.WINDOW_CLOSED:
        break
    if events == 'Gerar':
        preencherDados()
    elif events == "_Enter":
        preencherDados()
